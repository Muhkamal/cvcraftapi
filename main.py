"""
CVCraft API — Production FastAPI Backend
────────────────────────────────────────
Routes:
  Auth      POST /auth/register · POST /auth/login · GET /auth/me
  Resumes   GET/POST /resumes · GET/PUT/DELETE /resumes/{id}
  Export    POST /export/pdf · POST /export/docx
  Payments  POST /payments/initialize · POST /payments/verify · POST /payments/webhook
  AI        POST /ai/bullets · POST /ai/summary · POST /ai/skills · POST /ai/match
  User      GET /user/subscription · PUT /user/subscription
"""

import os
import hashlib
import hmac
import json
import httpx
import logging
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from typing import Optional

from fastapi import FastAPI, HTTPException, Depends, Header, BackgroundTasks, Request, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, EmailStr
from supabase import create_client, Client
from anthropic import AsyncAnthropic
import io

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("cvcraftapi")

# ─── ENV ──────────────────────────────────────────────────────────────────────
SUPABASE_URL       = os.environ["SUPABASE_URL"]
SUPABASE_KEY       = os.environ["SUPABASE_SERVICE_KEY"]       # service role key (server only)
SUPABASE_ANON_KEY  = os.environ["SUPABASE_ANON_KEY"]
ANTHROPIC_KEY      = os.environ["ANTHROPIC_API_KEY"]
PAYSTACK_SECRET    = os.environ["PAYSTACK_SECRET_KEY"]
PAYSTACK_WEBHOOK_SECRET = os.environ["PAYSTACK_WEBHOOK_SECRET"]
FRONTEND_URL       = os.environ.get("FRONTEND_URL", "http://localhost:5173")

# ─── CLIENTS ──────────────────────────────────────────────────────────────────
supabase: Client   = create_client(SUPABASE_URL, SUPABASE_KEY)
claude             = AsyncAnthropic(api_key=ANTHROPIC_KEY)

# ─── APP SETUP ────────────────────────────────────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    log.info("CVCraft API starting up")
    yield
    log.info("CVCraft API shutting down")

app = FastAPI(
    title="CVCraft API",
    version="1.0.0",
    description="AI-powered resume builder backend",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[FRONTEND_URL, "http://localhost:5173", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── MODELS ───────────────────────────────────────────────────────────────────
class RegisterRequest(BaseModel):
    email: EmailStr
    password: str
    full_name: Optional[str] = ""

class LoginRequest(BaseModel):
    email: EmailStr
    password: str

class ResumeData(BaseModel):
    title: Optional[str] = "Untitled Resume"
    template: Optional[str] = "executive"
    personal_info: dict = {}
    experience: list = []
    education: list = []
    skills: list = []
    certifications: list = []
    job_description: Optional[str] = ""

class ExportRequest(BaseModel):
    resume_id: Optional[str] = None
    resume_data: Optional[dict] = None   # inline data (no save needed)
    format: str = "pdf"                  # pdf | docx

class BulletsRequest(BaseModel):
    job_title: str
    company: Optional[str] = ""
    job_description: Optional[str] = ""
    count: int = 5

class SummaryRequest(BaseModel):
    job_title: str
    experience: list = []
    skills: list = []

class SkillsRequest(BaseModel):
    job_title: str
    count: int = 12

class MatchRequest(BaseModel):
    resume_data: dict
    job_description: str

class InitPaymentRequest(BaseModel):
    plan: str      # pro | lifetime
    email: EmailStr

class VerifyPaymentRequest(BaseModel):
    reference: str
    plan: str

# ─── AUTH HELPERS ─────────────────────────────────────────────────────────────
async def get_current_user(authorization: str = Header(None)) -> dict:
    """Verify Supabase JWT and return user dict."""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid Authorization header")
    token = authorization.split(" ", 1)[1]
    try:
        # Supabase verifies the JWT and returns the user
        resp = supabase.auth.get_user(token)
        if not resp or not resp.user:
            raise HTTPException(status_code=401, detail="Invalid token")
        return {"id": resp.user.id, "email": resp.user.email}
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"Auth error: {str(e)}")

async def get_user_plan(user_id: str) -> str:
    """Returns 'free' | 'pro' | 'lifetime'"""
    try:
        r = supabase.table("subscriptions").select("plan").eq("user_id", user_id).single().execute()
        return r.data.get("plan", "free") if r.data else "free"
    except Exception:
        return "free"

def require_plan(minimum: str):
    """Dependency factory: require minimum plan level."""
    order = {"free": 0, "pro": 1, "lifetime": 2}
    async def _check(user: dict = Depends(get_current_user)):
        plan = await get_user_plan(user["id"])
        if order.get(plan, 0) < order[minimum]:
            raise HTTPException(
                status_code=403,
                detail=f"This feature requires {minimum} plan. You are on {plan}."
            )
        return user
    return _check

# ─── HEALTH ───────────────────────────────────────────────────────────────────
@app.get("/", tags=["Health"])
async def root():
    return {"status": "ok", "service": "CVCraft API", "version": "1.0.0"}

@app.get("/health", tags=["Health"])
async def health():
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}

# ─── AUTH ROUTES ──────────────────────────────────────────────────────────────
@app.post("/auth/register", tags=["Auth"])
async def register(body: RegisterRequest):
    try:
        resp = supabase.auth.sign_up({
            "email": body.email,
            "password": body.password,
            "options": {"data": {"full_name": body.full_name}}
        })
        if resp.user:
            # Create free subscription record
            supabase.table("subscriptions").upsert({
                "user_id": resp.user.id,
                "plan": "free",
                "created_at": datetime.now(timezone.utc).isoformat(),
            }).execute()
        return {
            "user": {"id": resp.user.id, "email": resp.user.email},
            "session": {
                "access_token": resp.session.access_token if resp.session else None,
                "refresh_token": resp.session.refresh_token if resp.session else None,
            }
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/auth/login", tags=["Auth"])
async def login(body: LoginRequest):
    try:
        resp = supabase.auth.sign_in_with_password({
            "email": body.email,
            "password": body.password,
        })
        plan = await get_user_plan(resp.user.id)
        return {
            "user": {"id": resp.user.id, "email": resp.user.email, "plan": plan},
            "session": {
                "access_token": resp.session.access_token,
                "refresh_token": resp.session.refresh_token,
            }
        }
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid email or password")

@app.get("/auth/me", tags=["Auth"])
async def me(user: dict = Depends(get_current_user)):
    plan = await get_user_plan(user["id"])
    return {"id": user["id"], "email": user["email"], "plan": plan}

# ─── RESUME ROUTES ────────────────────────────────────────────────────────────
@app.get("/resumes", tags=["Resumes"])
async def list_resumes(user: dict = Depends(get_current_user)):
    plan = await get_user_plan(user["id"])
    r = supabase.table("resumes").select("id,title,template,updated_at,created_at").eq("user_id", user["id"]).order("updated_at", desc=True).execute()
    resumes = r.data or []
    # Free plan: cap at 1
    if plan == "free":
        resumes = resumes[:1]
    return {"resumes": resumes, "count": len(resumes), "plan": plan}

@app.post("/resumes", tags=["Resumes"], status_code=201)
async def create_resume(body: ResumeData, user: dict = Depends(get_current_user)):
    plan = await get_user_plan(user["id"])
    if plan == "free":
        existing = supabase.table("resumes").select("id").eq("user_id", user["id"]).execute()
        if len(existing.data or []) >= 1:
            raise HTTPException(status_code=403, detail="Free plan allows 1 resume. Upgrade to Pro for unlimited.")
    r = supabase.table("resumes").insert({
        "user_id": user["id"],
        "title": body.title,
        "template": body.template,
        "personal_info": body.personal_info,
        "experience": body.experience,
        "education": body.education,
        "skills": body.skills,
        "certifications": body.certifications,
        "job_description": body.job_description,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }).execute()
    return {"resume": r.data[0], "message": "Resume created"}

@app.get("/resumes/{resume_id}", tags=["Resumes"])
async def get_resume(resume_id: str, user: dict = Depends(get_current_user)):
    r = supabase.table("resumes").select("*").eq("id", resume_id).eq("user_id", user["id"]).single().execute()
    if not r.data:
        raise HTTPException(status_code=404, detail="Resume not found")
    return r.data

@app.put("/resumes/{resume_id}", tags=["Resumes"])
async def update_resume(resume_id: str, body: ResumeData, user: dict = Depends(get_current_user)):
    # Verify ownership
    existing = supabase.table("resumes").select("id").eq("id", resume_id).eq("user_id", user["id"]).single().execute()
    if not existing.data:
        raise HTTPException(status_code=404, detail="Resume not found")
    r = supabase.table("resumes").update({
        "title": body.title,
        "template": body.template,
        "personal_info": body.personal_info,
        "experience": body.experience,
        "education": body.education,
        "skills": body.skills,
        "certifications": body.certifications,
        "job_description": body.job_description,
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }).eq("id", resume_id).eq("user_id", user["id"]).execute()
    return {"resume": r.data[0], "message": "Resume updated"}

@app.delete("/resumes/{resume_id}", tags=["Resumes"], status_code=204)
async def delete_resume(resume_id: str, user: dict = Depends(get_current_user)):
    supabase.table("resumes").delete().eq("id", resume_id).eq("user_id", user["id"]).execute()
    return None

# ─── EXPORT ROUTES ────────────────────────────────────────────────────────────
def build_resume_html(resume: dict, template: str = "executive") -> str:
    """Build a clean, print-ready HTML for PDF generation."""
    p = resume.get("personal_info", {})
    experience = resume.get("experience", [])
    education = resume.get("education", [])
    skills = resume.get("skills", [])
    certifications = resume.get("certifications", [])

    # Template color map
    colors = {
        "executive": {"accent": "#4f46e5", "head_bg": "#fff",    "head_color": "#1a1a2e", "border": "#1a1a2e"},
        "minimal":   {"accent": "#999",    "head_bg": "#fff",    "head_color": "#111",    "border": "#eee"},
        "corporate": {"accent": "#1c2b4a", "head_bg": "#fff",    "head_color": "#1c2b4a", "border": "#1c2b4a"},
        "modern":    {"accent": "#4f46e5", "head_bg": "#0f0f1a", "head_color": "#fff",    "border": "#4f46e5"},
        "classic":   {"accent": "#222",    "head_bg": "#fff",    "head_color": "#111",    "border": "#222"},
        "tech":      {"accent": "#58a6ff", "head_bg": "#0d1117", "head_color": "#e6edf3", "border": "#21262d"},
        "elegant":   {"accent": "#c4a882", "head_bg": "#fff",    "head_color": "#1a1208", "border": "#d4c4a8"},
        "bold":      {"accent": "#e63946", "head_bg": "#1a1a2e", "head_color": "#fff",    "border": "#1a1a2e"},
    }
    c = colors.get(template, colors["classic"])

    def exp_html(e):
        bullets_html = ""
        if e.get("bullets"):
            bullets_html = "<ul style='margin:6px 0 0 16px;padding:0'>" + "".join(f"<li style='font-size:12px;color:#333;line-height:1.65;margin-bottom:2px'>{b}</li>" for b in e["bullets"]) + "</ul>"
        elif e.get("description"):
            bullets_html = f"<p style='font-size:12px;color:#333;line-height:1.7;margin:6px 0 0'>{e['description']}</p>"
        date_str = f"{e.get('startDate','')} – {'Present' if e.get('current') else e.get('endDate','')}"
        return f"""
        <div style='margin-bottom:14px'>
          <div style='display:flex;justify-content:space-between;align-items:baseline'>
            <strong style='font-size:13px;color:#111'>{e.get('jobTitle','')}</strong>
            <span style='font-size:11px;color:#888'>{date_str}</span>
          </div>
          <div style='font-size:12px;color:{c["accent"]};font-weight:600;margin-top:1px'>{e.get('company','')}{' · ' if e.get('company') and e.get('location') else ''}{e.get('location','')}</div>
          {bullets_html}
        </div>"""

    def edu_html(e):
        date_str = f"{e.get('startDate','')} – {e.get('endDate','')}"
        return f"""
        <div style='margin-bottom:12px'>
          <div style='display:flex;justify-content:space-between;align-items:baseline'>
            <strong style='font-size:13px;color:#111'>{e.get('degree','')}</strong>
            <span style='font-size:11px;color:#888'>{date_str}</span>
          </div>
          <div style='font-size:12px;color:{c["accent"]};margin-top:1px'>{e.get('school','')}</div>
          {f"<p style='font-size:12px;color:#555;margin-top:4px'>{e['description']}</p>" if e.get('description') else ''}
        </div>"""

    def cert_html(cert):
        return f"""
        <div style='margin-bottom:10px;display:flex;justify-content:space-between'>
          <div>
            <strong style='font-size:13px;color:#111'>{cert.get('name','')}</strong>
            <div style='font-size:12px;color:{c["accent"]}'>{cert.get('issuer','')}</div>
          </div>
          <span style='font-size:11px;color:#888'>{cert.get('date','')}</span>
        </div>"""

    skills_html = "".join(f"<span style='display:inline-block;font-size:12px;background:#f0f0f8;color:#333;padding:4px 12px;border-radius:3px;margin:3px;border:1px solid #ddd'>{s}</span>" for s in skills)
    contact_parts = [x for x in [p.get("email"), p.get("phone"), p.get("location"), p.get("linkedin"), p.get("website")] if x]
    contact_html = " &nbsp;·&nbsp; ".join(contact_parts)

    head_style = f"background:{c['head_bg']};color:{c['head_color']};padding:36px 48px 24px;border-bottom:2px solid {c['border']}"
    if template in ("modern", "tech", "bold"):
        head_style = f"background:{c['head_bg']};color:{c['head_color']};padding:36px 48px 24px"

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>{p.get('fullName','Resume')}</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&display=swap');
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{font-family:'DM Sans',Helvetica,Arial,sans-serif;color:#111;font-size:14px;-webkit-print-color-adjust:exact;print-color-adjust:exact}}
  @page{{size:A4;margin:0}}
  .page{{width:210mm;min-height:297mm;background:#fff}}
  .section-title{{font-size:10px;font-weight:700;text-transform:uppercase;letter-spacing:2px;color:{c['accent']};margin:20px 0 10px;padding-bottom:5px;border-bottom:1px solid #e8e8ee}}
</style>
</head>
<body>
<div class="page">
  <div style="{head_style}">
    <div style="font-size:28px;font-weight:700;letter-spacing:-0.5px">{p.get('fullName','Your Name')}</div>
    <div style="font-size:13px;margin-top:4px;opacity:.75">{p.get('jobTitle','')}</div>
    <div style="font-size:11px;margin-top:10px;opacity:.7">{contact_html}</div>
  </div>
  <div style="padding:28px 48px 40px">
    {f'<div><div class="section-title">Summary</div><p style="font-size:13px;color:#333;line-height:1.8">{p["summary"]}</p></div>' if p.get("summary") else ""}
    {f'<div><div class="section-title">Work Experience</div>{"".join(exp_html(e) for e in experience)}</div>' if experience else ""}
    {f'<div><div class="section-title">Education</div>{"".join(edu_html(e) for e in education)}</div>' if education else ""}
    {f'<div><div class="section-title">Skills</div><div style="margin-bottom:8px">{skills_html}</div></div>' if skills else ""}
    {f'<div><div class="section-title">Certifications</div>{"".join(cert_html(c) for c in certifications)}</div>' if certifications else ""}
  </div>
</div>
</body>
</html>"""

async def html_to_pdf(html: str) -> bytes:
    """Convert HTML to PDF using WeasyPrint."""
    try:
        from weasyprint import HTML, CSS
        pdf_bytes = HTML(string=html, base_url=None).write_pdf(
            stylesheets=[CSS(string="@page { size: A4; margin: 0; }")]
        )
        return pdf_bytes
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="WeasyPrint not installed. Run: pip install weasyprint"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

@app.post("/export/pdf", tags=["Export"])
async def export_pdf(body: ExportRequest, user: dict = Depends(get_current_user)):
    plan = await get_user_plan(user["id"])

    # Free plan: check daily download count
    if plan == "free":
        today = datetime.now(timezone.utc).date().isoformat()
        count_r = supabase.table("download_logs")\
            .select("id", count="exact")\
            .eq("user_id", user["id"])\
            .eq("date", today)\
            .execute()
        if (count_r.count or 0) >= 1:
            raise HTTPException(
                status_code=403,
                detail="Free plan allows 1 PDF download per day. Upgrade to Pro for unlimited."
            )

    # Get resume data
    if body.resume_id:
        r = supabase.table("resumes").select("*").eq("id", body.resume_id).eq("user_id", user["id"]).single().execute()
        if not r.data:
            raise HTTPException(status_code=404, detail="Resume not found")
        resume = r.data
        template = resume.get("template", "classic")
    elif body.resume_data:
        resume = body.resume_data
        template = resume.get("template", "classic")
    else:
        raise HTTPException(status_code=400, detail="Provide resume_id or resume_data")

    html = build_resume_html(resume, template)
    pdf_bytes = await html_to_pdf(html)

    # Log download (for free plan rate limiting)
    if plan == "free":
        supabase.table("download_logs").insert({
            "user_id": user["id"],
            "date": datetime.now(timezone.utc).date().isoformat(),
            "format": "pdf",
        }).execute()

    name = resume.get("personal_info", {}).get("fullName", "resume").replace(" ", "_")
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{name}_resume.pdf"'}
    )

@app.post("/export/docx", tags=["Export"])
async def export_docx(body: ExportRequest, user: dict = Depends(require_plan("pro"))):
    """DOCX export — Pro and above only."""
    if body.resume_id:
        r = supabase.table("resumes").select("*").eq("id", body.resume_id).eq("user_id", user["id"]).single().execute()
        if not r.data:
            raise HTTPException(status_code=404, detail="Resume not found")
        resume = r.data
    elif body.resume_data:
        resume = body.resume_data
    else:
        raise HTTPException(status_code=400, detail="Provide resume_id or resume_data")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p = resume.get("personal_info", {})
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        def add_heading(text, size=11, bold=True, color=(79, 70, 229)):
            p_obj = doc.add_paragraph()
            p_obj.paragraph_format.space_before = Pt(14)
            p_obj.paragraph_format.space_after = Pt(4)
            run = p_obj.add_run(text.upper())
            run.bold = bold
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(*color)
            p_obj.paragraph_format.border_bottom = True
            return p_obj

        def add_entry_header(title, subtitle, date):
            p_obj = doc.add_paragraph()
            p_obj.paragraph_format.space_before = Pt(8)
            p_obj.paragraph_format.space_after = Pt(1)
            r1 = p_obj.add_run(title)
            r1.bold = True
            r1.font.size = Pt(11)
            if date:
                r2 = p_obj.add_run(f"  {date}")
                r2.font.size = Pt(10)
                r2.font.color.rgb = RGBColor(100, 100, 120)
            if subtitle:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(2)
                r = p2.add_run(subtitle)
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(79, 70, 229)
                r.italic = True

        # Name & contact
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_r = name_p.add_run(p.get("fullName", "Your Name"))
        name_r.bold = True
        name_r.font.size = Pt(22)

        if p.get("jobTitle"):
            title_p = doc.add_paragraph()
            t_r = title_p.add_run(p["jobTitle"])
            t_r.font.size = Pt(12)
            t_r.font.color.rgb = RGBColor(100, 100, 120)

        contact_parts = [x for x in [p.get("email"), p.get("phone"), p.get("location"), p.get("linkedin"), p.get("website")] if x]
        if contact_parts:
            c_p = doc.add_paragraph(" | ".join(contact_parts))
            c_p.runs[0].font.size = Pt(9)
            c_p.runs[0].font.color.rgb = RGBColor(120, 120, 140)

        doc.add_paragraph()  # spacer

        if p.get("summary"):
            add_heading("Professional Summary")
            s_p = doc.add_paragraph(p["summary"])
            s_p.runs[0].font.size = Pt(11)

        if resume.get("experience"):
            add_heading("Work Experience")
            for exp in resume["experience"]:
                date_str = f"{exp.get('startDate','')} – {'Present' if exp.get('current') else exp.get('endDate','')}"
                company = exp.get("company","") + (" · " if exp.get("company") and exp.get("location") else "") + exp.get("location","")
                add_entry_header(exp.get("jobTitle",""), company, date_str)
                for bullet in (exp.get("bullets") or []):
                    b_p = doc.add_paragraph(bullet, style="List Bullet")
                    b_p.runs[0].font.size = Pt(10)
                if exp.get("description") and not exp.get("bullets"):
                    d_p = doc.add_paragraph(exp["description"])
                    d_p.runs[0].font.size = Pt(10)

        if resume.get("education"):
            add_heading("Education")
            for edu in resume["education"]:
                date_str = f"{edu.get('startDate','')} – {edu.get('endDate','')}"
                add_entry_header(edu.get("degree",""), edu.get("school",""), date_str)
                if edu.get("description"):
                    d_p = doc.add_paragraph(edu["description"])
                    d_p.runs[0].font.size = Pt(10)

        if resume.get("skills"):
            add_heading("Skills")
            s_p = doc.add_paragraph(" · ".join(resume["skills"]))
            s_p.runs[0].font.size = Pt(10)

        if resume.get("certifications"):
            add_heading("Certifications")
            for cert in resume["certifications"]:
                add_entry_header(cert.get("name",""), cert.get("issuer",""), cert.get("date",""))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        name = p.get("fullName", "resume").replace(" ", "_")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{name}_resume.docx"'}
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        log.error(f"DOCX export failed: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {str(e)}")

# ─── AI ROUTES ────────────────────────────────────────────────────────────────
@app.post("/ai/bullets", tags=["AI"])
async def ai_bullets(body: BulletsRequest, user: dict = Depends(get_current_user)):
    plan = await get_user_plan(user["id"])
    count = min(body.count, 8 if plan != "free" else 5)

    job_ctx = f"\nTarget job posting:\n{body.job_description[:800]}" if body.job_description else ""
    prompt = f"""Generate {count} powerful, ATS-optimized resume bullet points for:
Job Title: {body.job_title}
Company: {body.company or 'Not specified'}{job_ctx}

Rules:
- Start each bullet with a strong action verb (Led, Built, Increased, Reduced, etc.)
- Include quantified results where possible (%, $, time saved, team size)
- Keep each bullet under 20 words
- Make bullets specific and impactful, not generic
- Tailor to the job posting if provided

Return ONLY a valid JSON array of {count} strings. No explanation."""

    try:
        resp = await claude.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=600,
            system="You are an expert resume writer and career coach. Return only valid JSON arrays.",
            messages=[{"role": "user", "content": prompt}]
        )
        raw = resp.content[0].text.strip().replace("```json", "").replace("```", "").strip()
        bullets = json.loads(raw)
        return {"bullets": bullets if isinstance(bullets, list) else [], "plan": plan}
    except Exception as e:
        log.error(f"AI bullets error: {e}")
        raise HTTPException(status_code=500, detail="AI generation failed")

@app.post("/ai/summary", tags=["AI"])
async def ai_summary(body: SummaryRequest, user: dict = Depends(get_current_user)):
    exp_list = ", ".join(f"{e.get('jobTitle','')} at {e.get('company','')}" for e in body.experience[:4])
    skills_list = ", ".join(body.skills[:15])
    prompt = f"""Write a compelling 2–3 sentence professional resume summary for:
Job Title: {body.job_title}
Experience: {exp_list or 'Not specified'}
Key Skills: {skills_list or 'Not specified'}

Rules:
- Be specific, not generic ("Experienced professional" is bad)
- Mention years of experience if inferable
- Lead with a strong identifier
- End with value proposition
- Under 60 words total

Return ONLY the summary text. No labels, no quotes."""

    try:
        resp = await claude.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=200,
            system="You are an expert resume writer. Write powerful, specific summaries.",
            messages=[{"role": "user", "content": prompt}]
        )
        return {"summary": resp.content[0].text.strip()}
    except Exception as e:
        raise HTTPException(status_code=500, detail="AI generation failed")

@app.post("/ai/skills", tags=["AI"])
async def ai_skills(body: SkillsRequest, user: dict = Depends(get_current_user)):
    prompt = f"""List {body.count} relevant, in-demand resume skills for a {body.job_title}.
Mix technical and soft skills. Be specific (e.g., "React.js" not "JavaScript frameworks").
Return ONLY a valid JSON array of strings."""

    try:
        resp = await claude.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=300,
            system="Return only valid JSON arrays.",
            messages=[{"role": "user", "content": prompt}]
        )
        raw = resp.content[0].text.strip().replace("```json", "").replace("```", "").strip()
        skills = json.loads(raw)
        return {"skills": skills if isinstance(skills, list) else []}
    except Exception as e:
        raise HTTPException(status_code=500, detail="AI generation failed")

@app.post("/ai/match", tags=["AI"])
async def ai_match(body: MatchRequest, user: dict = Depends(require_plan("pro"))):
    """Job match analysis — Pro+ only."""
    p = body.resume_data.get("personal_info", {})
    experience = body.resume_data.get("experience", [])
    skills = body.resume_data.get("skills", [])

    resume_text = "\n".join([
        p.get("summary", ""),
        *[f"{e.get('jobTitle','')} at {e.get('company','')} — {e.get('description','')} {' '.join(e.get('bullets', []))}" for e in experience[:5]],
        "Skills: " + ", ".join(skills),
    ])

    prompt = f"""Analyze how well this resume matches the job description.

RESUME:
{resume_text[:1200]}

JOB DESCRIPTION:
{body.job_description[:1200]}

Return ONLY valid JSON with this exact structure:
{{
  "score": <integer 0-100>,
  "label": "<Strong Match|Good Match|Fair Match|Needs Work>",
  "missing_keywords": ["keyword1", "keyword2"],
  "tips": ["Specific tip 1", "Specific tip 2", "Specific tip 3", "Specific tip 4"],
  "strengths": ["Strength 1", "Strength 2"]
}}"""

    try:
        resp = await claude.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=500,
            system="You are an ATS expert. Return only valid JSON.",
            messages=[{"role": "user", "content": prompt}]
        )
        raw = resp.content[0].text.strip().replace("```json", "").replace("```", "").strip()
        result = json.loads(raw)
        return result
    except Exception as e:
        log.error(f"AI match error: {e}")
        raise HTTPException(status_code=500, detail="AI analysis failed")

# ─── PAYMENT ROUTES ───────────────────────────────────────────────────────────
PLAN_AMOUNTS = {
    "pro":      450000,   # ₦4,500 in kobo
    "lifetime": 1990000,  # ₦19,900 in kobo
}

@app.post("/payments/initialize", tags=["Payments"])
async def initialize_payment(body: InitPaymentRequest, user: dict = Depends(get_current_user)):
    """Initialize a Paystack transaction and return authorization URL."""
    if body.plan not in PLAN_AMOUNTS:
        raise HTTPException(status_code=400, detail="Invalid plan. Must be 'pro' or 'lifetime'")

    amount = PLAN_AMOUNTS[body.plan]
    reference = f"CVCRAFT_{body.plan.upper()}_{user['id'][:8]}_{int(datetime.now().timestamp())}"

    async with httpx.AsyncClient() as client:
        resp = await client.post(
            "https://api.paystack.co/transaction/initialize",
            headers={
                "Authorization": f"Bearer {PAYSTACK_SECRET}",
                "Content-Type": "application/json",
            },
            json={
                "email": body.email,
                "amount": amount,
                "currency": "NGN",
                "reference": reference,
                "metadata": {
                    "user_id": user["id"],
                    "plan": body.plan,
                    "custom_fields": [
                        {"display_name": "Plan", "variable_name": "plan", "value": body.plan},
                        {"display_name": "User ID", "variable_name": "user_id", "value": user["id"]},
                    ]
                },
                "callback_url": f"{FRONTEND_URL}/payment/callback",
            }
        )

    data = resp.json()
    if not data.get("status"):
        raise HTTPException(status_code=400, detail=data.get("message", "Paystack initialization failed"))

    # Log pending transaction
    supabase.table("payment_logs").insert({
        "user_id": user["id"],
        "reference": reference,
        "plan": body.plan,
        "amount": amount,
        "status": "pending",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }).execute()

    return {
        "authorization_url": data["data"]["authorization_url"],
        "access_code": data["data"]["access_code"],
        "reference": reference,
    }

@app.post("/payments/verify", tags=["Payments"])
async def verify_payment(body: VerifyPaymentRequest, user: dict = Depends(get_current_user)):
    """Verify a Paystack transaction by reference and upgrade user plan."""
    async with httpx.AsyncClient() as client:
        resp = await client.get(
            f"https://api.paystack.co/transaction/verify/{body.reference}",
            headers={"Authorization": f"Bearer {PAYSTACK_SECRET}"},
        )

    data = resp.json()
    if not data.get("status") or data["data"]["status"] != "success":
        raise HTTPException(status_code=400, detail="Payment not successful")

    tx = data["data"]
    metadata = tx.get("metadata", {})
    paid_user_id = metadata.get("user_id", user["id"])

    # Security: ensure the reference belongs to this user
    if paid_user_id != user["id"]:
        raise HTTPException(status_code=403, detail="Reference does not belong to this account")

    plan = metadata.get("plan", body.plan)

    # Upsert subscription
    supabase.table("subscriptions").upsert({
        "user_id": user["id"],
        "plan": plan,
        "paystack_reference": body.reference,
        "amount_paid": tx["amount"],
        "activated_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }, on_conflict="user_id").execute()

    # Update payment log
    supabase.table("payment_logs").update({
        "status": "success",
        "verified_at": datetime.now(timezone.utc).isoformat(),
    }).eq("reference", body.reference).execute()

    log.info(f"Plan upgraded: user={user['id']} plan={plan} ref={body.reference}")
    return {"success": True, "plan": plan, "message": f"Successfully upgraded to {plan.capitalize()}!"}

@app.post("/payments/webhook", tags=["Payments"])
async def paystack_webhook(request: Request, background_tasks: BackgroundTasks):
    """
    Paystack webhook endpoint.
    Verifies HMAC signature then processes charge.success events.
    Set this URL in your Paystack dashboard → Settings → Webhooks.
    """
    body_bytes = await request.body()
    signature = request.headers.get("x-paystack-signature", "")

    # Verify webhook signature
    expected = hmac.new(
        PAYSTACK_WEBHOOK_SECRET.encode("utf-8"),
        body_bytes,
        hashlib.sha512
    ).hexdigest()

    if not hmac.compare_digest(expected, signature):
        log.warning("Webhook signature mismatch — rejecting")
        raise HTTPException(status_code=400, detail="Invalid webhook signature")

    payload = json.loads(body_bytes)
    event = payload.get("event")
    tx_data = payload.get("data", {})

    log.info(f"Webhook received: event={event} ref={tx_data.get('reference')}")

    if event == "charge.success":
        background_tasks.add_task(_handle_successful_charge, tx_data)

    return {"received": True}

async def _handle_successful_charge(tx: dict):
    """Background task: activate plan after confirmed webhook."""
    try:
        metadata = tx.get("metadata", {})
        user_id = metadata.get("user_id")
        plan = metadata.get("plan")
        reference = tx.get("reference")

        if not user_id or not plan:
            log.error(f"Webhook missing metadata: {metadata}")
            return

        supabase.table("subscriptions").upsert({
            "user_id": user_id,
            "plan": plan,
            "paystack_reference": reference,
            "amount_paid": tx.get("amount"),
            "activated_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }, on_conflict="user_id").execute()

        supabase.table("payment_logs").update({
            "status": "webhook_confirmed",
            "verified_at": datetime.now(timezone.utc).isoformat(),
        }).eq("reference", reference).execute()

        log.info(f"Webhook processed: user={user_id} plan={plan} ref={reference}")
    except Exception as e:
        log.error(f"Webhook processing error: {e}")

# ─── USER ROUTES ──────────────────────────────────────────────────────────────
@app.get("/user/subscription", tags=["User"])
async def get_subscription(user: dict = Depends(get_current_user)):
    plan = await get_user_plan(user["id"])
    r = supabase.table("subscriptions").select("*").eq("user_id", user["id"]).single().execute()
    return {"plan": plan, "details": r.data}

@app.get("/user/stats", tags=["User"])
async def get_stats(user: dict = Depends(get_current_user)):
    plan = await get_user_plan(user["id"])
    resume_count = supabase.table("resumes").select("id", count="exact").eq("user_id", user["id"]).execute()
    today = datetime.now(timezone.utc).date().isoformat()
    downloads_today = supabase.table("download_logs").select("id", count="exact").eq("user_id", user["id"]).eq("date", today).execute()
    return {
        "plan": plan,
        "resume_count": resume_count.count or 0,
        "downloads_today": downloads_today.count or 0,
        "limits": {
            "max_resumes": 1 if plan == "free" else None,
            "downloads_per_day": 1 if plan == "free" else None,
        }
    }
